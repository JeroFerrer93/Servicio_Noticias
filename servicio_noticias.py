import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import nltk
nltk.download("punkt")  # ✅ necesario para sumarizador


# 🔧 Limpiar caracteres no válidos
def limpiar_texto(texto):
    return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", texto).strip()

# 🧠 Resumen automático
def resumir_texto(texto, oraciones=3):
    try:
        parser = PlaintextParser.from_string(texto, Tokenizer("spanish"))
        resumen = LsaSummarizer()(parser.document, oraciones)
        resultado = " ".join(str(oracion) for oracion in resumen)
        return resultado if resultado.strip() else "⚠️ No se pudo generar el resumen."
    except Exception as e:
        print(f"⚠️ Error al resumir: {e}")
        return texto[:800] + "..." if texto else "⚠️ No se pudo generar el resumen."


# 🌐 Mundo Gremial
def obtener_noticias_mundo_gremial():
    url = "https://mundogremial.com/"
    respuesta = requests.get(url)
    if respuesta.status_code != 200:
        print(f"❌ Error al acceder a {url}")
        return []

    sopa = BeautifulSoup(respuesta.text, "html.parser")
    noticias = []

    for articulo in sopa.find_all("h2", class_="mvp-stand-title"):
        titulo = articulo.get_text(strip=True)
        enlace = articulo.find_parent("a")["href"] if articulo.find_parent("a") else None
        if titulo and enlace:
            noticias.append({"titulo": titulo, "fuente": "Mundo Gremial", "enlace": enlace})

    print(f"🔍 Se encontraron {len(noticias)} noticias en Mundo Gremial.")
    return noticias

# 🌐 Ámbito
def obtener_noticias_ambito():
    url = "https://www.ambito.com/"
    headers = {"User-Agent": "Mozilla/5.0"}
    noticias = []

    try:
        respuesta = requests.get(url, headers=headers)
        if respuesta.status_code != 200:
            print(f"❌ Error al acceder a {url}")
            return []

        sopa = BeautifulSoup(respuesta.text, "html.parser")
        titulos = sopa.find_all("h2")[:15]  # 🔢 Limita a las primeras 15 noticias

        for h2 in titulos:
            link_tag = h2.find("a", href=True)
            if link_tag:
                titulo = h2.get_text(strip=True)
                href = link_tag["href"]
                if href and not href.startswith("http"):
                    href = "https://www.ambito.com" + href
                noticias.append({
                    "titulo": titulo,
                    "fuente": "Ámbito",
                    "enlace": href
                })

    except Exception as e:
        print(f"⚠️ Error procesando Ámbito: {e}")

    print(f"🔍 Se encontraron {len(noticias)} noticias en Ámbito.")
    return noticias

# 🌐 Cronista
def obtener_noticias_cronista():
    url = "https://www.cronista.com/economia-politica/"
    headers = {"User-Agent": "Mozilla/5.0"}
    noticias = []

    try:
        respuesta = requests.get(url, headers=headers)
        if respuesta.status_code != 200:
            print(f"❌ Error al acceder a {url}")
            return []

        sopa = BeautifulSoup(respuesta.text, "html.parser")

        # Mostrar todos los enlaces que encuentre
        links = sopa.select("a[href]")
        print(f"🔧 Total de <a> encontrados: {len(links)}")

        for a in links:
            href = a.get("href")
            titulo = a.get_text(strip=True)
            if href and "/economia-politica/" in href and len(titulo) > 20:
                if not href.startswith("http"):
                    href = "https://www.cronista.com" + href
                noticias.append({
                    "titulo": titulo,
                    "fuente": "El Cronista",
                    "enlace": href
                })

    except Exception as e:
        print(f"⚠️ Error procesando El Cronista: {e}")

    print(f"🔍 Se encontraron {len(noticias)} noticias en El Cronista.")
    return noticias


# 📥 Extraer contenido de cada nota
def obtener_contenido_noticia(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        respuesta = requests.get(url, headers=headers)
        if respuesta.status_code != 200:
            return "No se pudo acceder al contenido."

        sopa = BeautifulSoup(respuesta.text, "html.parser")

        if "mundogremial.com" in url:
            contenido_div = sopa.find("div", id="mvp-content-main")
        elif "ambito.com" in url:
            contenido_div = sopa.find("main")
        elif "cronista.com" in url:
            contenido_div = sopa.find("div", class_="content vsmcontent news")
        else:
            contenido_div = None

        if contenido_div:
            parrafos = contenido_div.find_all("p")
            texto = ' '.join(p.get_text(strip=True) for p in parrafos if p.get_text(strip=True))
            return texto if texto and len(texto) > 50 else "No se encontró contenido relevante."

    except Exception as e:
        print(f"❌ Error al obtener contenido de {url}: {e}")

    return "No se pudo extraer el contenido."

# 📄 Generar documento Word
def generar_word():
    doc = Document()
    doc.add_heading('Resumen de Noticias Gremiales y Económicas', level=1)

    noticias = []
    noticias.extend(obtener_noticias_mundo_gremial())
    noticias.extend(obtener_noticias_ambito())
    noticias.extend(obtener_noticias_cronista())

    # 🧹 Eliminar noticias duplicadas por título
    titulos_vistos = set()
    noticias_unicas = []
    for n in noticias:
        if n["titulo"] not in titulos_vistos:
            titulos_vistos.add(n["titulo"])
            noticias_unicas.append(n)
    noticias = noticias_unicas

    for noticia in noticias:
        doc.add_heading(limpiar_texto(noticia["titulo"]), level=2)
        doc.add_paragraph(f"📌 Fuente: {limpiar_texto(noticia['fuente'])}")
        doc.add_paragraph(f"🔗 Enlace: {limpiar_texto(noticia['enlace'])}")

        contenido = obtener_contenido_noticia(noticia["enlace"])
        resumen = resumir_texto(contenido)

        doc.add_paragraph(f"📝 Resumen: {limpiar_texto(resumen)}\n")
        doc.add_paragraph("📰────────────────────────────")

    fecha_actual = datetime.now().strftime("%Y-%m-%d_%H-%M")
    nombre_archivo = f"Resumen_Noticias_{fecha_actual}.docx"
    doc.save(nombre_archivo)
    print(f"✅ Documento Word generado: {nombre_archivo}")
    return nombre_archivo    

# 🚀 Ejecutar
if __name__ == "__main__":
    generar_word()
